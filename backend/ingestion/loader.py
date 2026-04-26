"""
Universal document loader.
Handles: PDF, DOCX, XLSX/CSV, HTML, plain text, and web URLs.
Returns a list of dicts: {content, page, metadata}
"""
import os
import io
import csv
import requests
from pathlib import Path
from typing import Literal
import structlog

logger = structlog.get_logger()

SupportedType = Literal["pdf", "docx", "xlsx", "csv", "html", "txt", "url"]


def detect_type(path: str) -> SupportedType:
    ext = Path(path).suffix.lower().lstrip(".")
    mapping = {
        "pdf": "pdf",
        "docx": "docx",
        "doc": "docx",
        "xlsx": "xlsx",
        "xls": "xlsx",
        "csv": "csv",
        "html": "html",
        "htm": "html",
        "txt": "txt",
        "md": "txt",
    }
    return mapping.get(ext, "txt")


def load_pdf(path: str) -> list[dict]:
    import fitz  # PyMuPDF
    doc = fitz.open(path)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text("text").strip()
        if text:
            pages.append({"content": text, "page": i + 1, "metadata": {"page": i + 1}})
    doc.close()
    return pages


def load_docx(path: str) -> list[dict]:
    from docx import Document
    doc = Document(path)
    chunks = []
    current = []
    for para in doc.paragraphs:
        if para.text.strip():
            current.append(para.text.strip())
        elif current:
            chunks.append({"content": "\n".join(current), "page": None, "metadata": {}})
            current = []
    if current:
        chunks.append({"content": "\n".join(current), "page": None, "metadata": {}})
    return chunks or [{"content": "\n".join(p.text for p in doc.paragraphs), "page": None, "metadata": {}}]


def load_xlsx(path: str) -> list[dict]:
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    pages = []
    for sheet in wb.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            row_text = "\t".join(str(c) for c in row if c is not None)
            if row_text.strip():
                rows.append(row_text)
        if rows:
            pages.append({
                "content": f"Sheet: {sheet.title}\n" + "\n".join(rows),
                "page": None,
                "metadata": {"sheet": sheet.title},
            })
    return pages


def load_csv(path: str) -> list[dict]:
    with open(path, newline="", encoding="utf-8-sig") as f:
        reader = csv.reader(f)
        rows = ["\t".join(row) for row in reader if any(row)]
    batch_size = 100
    chunks = []
    for i in range(0, len(rows), batch_size):
        chunks.append({
            "content": "\n".join(rows[i:i + batch_size]),
            "page": None,
            "metadata": {"rows": f"{i+1}-{min(i+batch_size, len(rows))}"},
        })
    return chunks


def load_html(path_or_content: str, is_content: bool = False) -> list[dict]:
    from bs4 import BeautifulSoup
    if is_content:
        html = path_or_content
    else:
        with open(path_or_content, encoding="utf-8") as f:
            html = f.read()
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    return [{"content": text, "page": None, "metadata": {}}]


def load_txt(path: str) -> list[dict]:
    with open(path, encoding="utf-8") as f:
        text = f.read()
    return [{"content": text, "page": None, "metadata": {}}]


def load_url(url: str) -> list[dict]:
    headers = {"User-Agent": "Mozilla/5.0 (RAG-Bot/1.0)"}
    resp = requests.get(url, headers=headers, timeout=15)
    resp.raise_for_status()
    content_type = resp.headers.get("content-type", "")
    if "pdf" in content_type:
        import fitz
        doc = fitz.open(stream=resp.content, filetype="pdf")
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text("text").strip()
            if text:
                pages.append({"content": text, "page": i + 1, "metadata": {"url": url, "page": i + 1}})
        return pages
    else:
        chunks = load_html(resp.text, is_content=True)
        for c in chunks:
            c["metadata"]["url"] = url
        return chunks


def load_document(path: str, file_type: SupportedType | None = None) -> list[dict]:
    """Main entry point. Returns list of page/section dicts."""
    # Check if it's a URL
    is_url_type = file_type == "url"
    is_url_path = path.startswith(("http://", "https://"))
    
    if is_url_type or is_url_path:
        url = path
        if not url.startswith(("http://", "https://")):
            url = f"https://{url}"
        logger.info("loading_url", url=url)
        return load_url(url) /url 

    ftype = file_type or detect_type(path)
    logger.info("loading_document", path=path, type=ftype)

    loaders = {
        "pdf": load_pdf,
        "docx": load_docx,
        "xlsx": load_xlsx,
        "csv": load_csv,
        "html": load_html,
        "txt": load_txt,
    }
    loader = loaders.get(ftype, load_txt)
    return loader(path)
